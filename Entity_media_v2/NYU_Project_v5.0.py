# -*- coding: utf-8 -*-
"""
Created on Wed May  6 21:47:25 2020
@author: Administrator
"""
#References
#https://www.learndatasci.com/tutorials/sentiment-analysis-reddit-headlines-pythons-nltk/

import requests
from requests import get
from requests.exceptions import RequestException
from bs4 import BeautifulSoup
from googlesearch import search
import nltk
import nltk
#https://stackoverflow.com/questions/36516697/not-able-to-install-nltk-data-on-django-app-on-elastic-beanstalk/36526192#36526192 for bringing in the #ntlk data on AWS beanstalj environment
try:
    nltk.download('punkt', download_dir='/opt/python/current/app')
    nltk.download('stopwords', download_dir='/opt/python/current/app')
    nltk.download('vader_lexicon', download_dir='/opt/python/current/app')
except:
    nltk.download('punkt')
    nltk.download('vader_lexicon')
    nltk.download('stopwords')
#nltk.download('punkt')
from nltk.corpus import stopwords #nltk.download('stopwords') if unavailable
from nltk.tokenize import word_tokenize # nltk.download('punkt')  if unavailable
from nltk.sentiment.vader import SentimentIntensityAnalyzer as SIA
import string
import pandas as pd
import wikipedia
from pandas.io.html import read_html
from docx import Document

# to search 
#query = "Osama Bin Laden"

class media_analyzer():
    

    def __init__(self,query,number_results):
        self.query = query
        self.number_of_results = number_results
 
    # function to extract links to google pages
    def link_extractor (self):
        self.html_link = []
        query = self.query
        numb = self.number_of_results
        for j in search(query, tld="com", num=numb, stop=numb, pause=4):
            self.html_link.append(j)
        return self.html_link
    
    # function to extract text from search pages
    def text_extractor (self,source_list): # takes in list of html links
        self.text_list = []
        for link in source_list:
            if requests.get(link).status_code == 200:
                soup = BeautifulSoup(requests.get(link).text, 'html.parser')
                page_text = soup.find_all('p')
                text = ""
                for element in page_text:
                    if element.text != "":
                        text += (element.text).encode("ascii", errors="ignore").decode(errors="ignore").lower()
                        text += " "
                self.text_list.append(text)
            else:
                pass
        return self.text_list # returns a list of teext parsed from searched html links
    
    # function to remove punctuation, remove stopwords, and tokenize string content to words
    @staticmethod
    def cleaner(review): # takes in string
        stop_words = set(stopwords.words('english'))
        word_list = word_tokenize(review)
        table = str.maketrans('', '', string.punctuation)
        stripped = [w.translate(table) for w in word_list]
        text = ""
        for word in stripped:
            if not word in stop_words:
                text += word
                text += " "
        return text  # retunrs string
    
    #function to generate text summary statistics
    @classmethod
    def text_summary(cls,text_list):# takes in a list
        wordcount = {}
        for i in text_list:
            strng = cls.cleaner(i.lower())  #takes string,returns string
            for word in strng.lower().split():
                if word not in wordcount:
                    wordcount[word] = 1
                else:
                    wordcount[word] += 1
        return sorted(wordcount.items(),key=lambda kv:(kv[1], kv[0]), reverse = True)
    
    @classmethod
    def TF_ML_ABC_Keywords(cls):
        ABC_keywords = []
        TF_keywords = []
        ML_keywords = []
        f = open("ABC_Keywords.txt", "r",errors = 'ignore')
        for x in f:
            ABC_keywords.append(((cls.cleaner(x)).lower()).strip())
        f = open("AML_Keywords.txt", "r",errors = 'ignore')
        for x in f:
            ML_keywords.append(((cls.cleaner(x)).lower()).strip())
        f = open("TF_Keywords.txt", "r",errors = 'ignore')
        for x in f:
            TF_keywords.append(((cls.cleaner(x)).lower()).strip())
        return ABC_keywords,TF_keywords,ML_keywords
    
    @classmethod
    def text_summary_keywords(cls,text_list,keyword_list):# takes in a list
        wordcount = {}
        for i in text_list:
            strng = cls.cleaner(i.lower())  #takes string,returns string
            for word in strng.lower().split():
                if word in keyword_list: 
                    if word not in wordcount:
                        wordcount[word] = 1
                    else:
                        wordcount[word] += 1
        return sorted(wordcount.items(),key=lambda kv:(kv[1], kv[0]), reverse = True)
    
    
    # function for sentiment analysis using NLTK's inbult sentiment analyzer
    @classmethod
    def polarity_calculator(cls,string_data): # takes in string
        sia = SIA()
        strng = cls.cleaner(string_data)
        pol_score = sia.polarity_scores((strng))
        pol_score['article'] = string_data
        return pol_score  #returns dict
    
    # funtion to generate the new score i.e good media or negative media coverage on the bases of search history 
    @classmethod
    def entity_score(cls,score_list):# takes in list
        df = pd.DataFrame.from_records(score_list)
        df['label'] = 0
        df.loc[df['compound'] > 0.2, 'label'] = 1
        df.loc[df['compound'] < -0.2, 'label'] = -1
        gk = df.groupby('label')
        score = len(gk.get_group(1)) - len(gk.get_group(-1))
        #entity = media_analyzer(self.query,self.number_results,self.entity_score)
        return score


class case_narrative():
    # ASSIGN THE VARIABLES FROM MEDIA ANALYZER CLASS OUTPUT
    def __init__(self,entity_name,final_score,out_1,keyword_list,page,ABC_list,ML_list,TF_list):
        self.entity_name = entity_name
        self.final_score = final_score
        self.out_1 = out_1
        self.keyword_list = keyword_list
        self.page = page

  # EXTRACT THE PROFILE INFORMATION FROM WIKIPEDIA
        wiki_search = wikipedia.search(entity_name,10)

        if final_score >0:
            sentiment = "Postitve"
            
        if final_score <=0:
            sentiment = "Negative"

        ML_keyword = 'Fraud'
        selection = wiki_search[1]
        selection_page = wikipedia.page(selection)
        selection_url = selection_page.url
        page = selection_url

# EXTRACT INFOBOX FROM WIKIPEDIA
        try:
            malformed = read_html(page, index_col=0, attrs={"class":"infobox"})
            malformed[0]
            Err = 'No'

        except ValueError:
            Err = 'yes'    
# CREATE THE CASE NARRATIVE WORD DOUCMENT AND PLACE THE PROFILE INFORMATION
        summ = (wikipedia.summary(selection, sentences = 2))
        doc = Document()
        doc.add_heading ('Case Narrative Template', 0)
        doc.add_heading ('Entity Name :' + entity_name, 2)
        doc.add_heading ('Introduction', 2)
        doc.add_paragraph(summ)

 # INFOBOX LOADED TO CASE NARRATIVE DOCUMENT

        if Err != 'yes':
            df = pd.DataFrame(malformed)
            t = doc.add_table(df.shape[0]+1, df.shape[1])
            # add the rest of the data frame
            for i in range(df.shape[0]):
                for j in range(df.shape[-1]):
                    t.cell(i+1,j).text = str(df.values[i,j])

        doc.add_heading ('Entity Media analyser', 2)
        doc.add_paragraph ('Sentiment :'+ sentiment)
        doc.add_paragraph ('Final Score :'+ str(final_score))
        doc.add_heading ('Top 5 Key words in the webpage',2)
  
   # TOP 5 KEY WORD LIST FROM MEDIA ANALYSER CLASS    
        for i in keyword_list:
            doc.add_paragraph (str(i))

        doc.add_heading ('Top 5 Money Laundering Key words in the webpage',2)
  
   # TOP 5 KEY WORD LIST FROM MEDIA ANALYSER CLASS    
        for i in ML_list:
            doc.add_paragraph (str(i))    

        doc.add_heading ('Top 5 Terrorist Financing Key words in the webpage',2)
  
   # TOP 5 KEY WORD LIST FROM MEDIA ANALYSER CLASS    
        for i in TF_list:
            doc.add_paragraph (str(i))             
  
        doc.add_heading ('Top 5 Bribery & Corruption Key words in the webpage',2)
  
   # TOP 5 KEY WORD LIST FROM MEDIA ANALYSER CLASS    
        for i in ABC_list:
            doc.add_paragraph (str(i)) 

        doc.add_paragraph ('Keywords releated to compliance: ' + ML_keyword)
        doc.add_heading ('Insights', 2)
        doc.add_paragraph('Based on the below web page search, Focal entity '+entity_name+' appear in ' +page+ ' web pages and overall sentiment is ' + sentiment + 'and sentiment final score is '+ str(final_score))
        doc.add_heading ('Sources', 2)

        for i in  out_1:
            doc.add_paragraph(i)

        doc.save('case_summary'+ entity_name +'.docx')
        
# FLASK / REACT APP LOADING....
import flask
app = flask.Flask(__name__)           # a Flask object

@app.route('/details', methods=['POST', 'GET'])
def entity_analysis():

    if  flask.request.method== 'GET':       
    
        entity_name = flask.request.args.get('ent')
        page = int(flask.request.args.get('page'))
        worddoc = flask.request.args.get('narrative')
        no_id = flask.request.args.get('no_id')  # from a GET (URL)
        media_analyzer(entity_name,page)
        
        if entity_name:
            page = page
            inst = media_analyzer(entity_name,page)
            out_1 = inst.link_extractor()
            out_2 = inst.text_extractor(out_1)
            out_3 = []
            for i in out_2:
                score = media_analyzer.polarity_calculator(i)
                out_3.append(score)
            final_score = media_analyzer.entity_score(out_3)
            
            if final_score >0:
                sentiment = "Postitve"
            if final_score <=0:
                sentiment = "Negative"

            keyword_list = (media_analyzer.text_summary(out_2)[:5])
            ABC_keywords,TF_keywords,ML_keywords = media_analyzer.TF_ML_ABC_Keywords()
            ABC_list = (media_analyzer.text_summary_keywords(out_2,ABC_keywords)[:5])
            TF_list = (media_analyzer.text_summary_keywords(out_2,TF_keywords)[:5])
            ML_list = (media_analyzer.text_summary_keywords(out_2,ML_keywords)[:5])

            
            if worddoc == 'Yes':
                inst1 = case_narrative(entity_name,final_score,out_1,keyword_list,page,TF_list,ML_list,ABC_list)
           
            msg = flask.render_template('Result_Template.html', obj = inst, score = final_score, word_s = keyword_list, final_sentiment = sentiment,word_TF = TF_list,word_ML= ML_list, word_ABC = ABC_list)
            
        elif no_id:
            msg = 'No Object Details.'

        else:
            
            raise ValueError('\nraised error:  no "name" or "no_name" params passed in request')

        return '<PRE>{}</PRE>'.format(msg)

if __name__ == '__main__':
    app.run(debug=True, port=5000)    # app starts serving in debug mode on port 5000